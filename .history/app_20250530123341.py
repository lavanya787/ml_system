from flask import Flask, request, jsonify, send_from_directory
import pandas as pd
import os
import uuid
import joblib
import json
from datetime import datetime

# File parsing
import docx
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from werkzeug.utils import secure_filename

# ML
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score

# NLP
import spacy
from transformers import pipeline

# PDF generation
from fpdf import FPDF

# --- Setup ---
app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ANALYSIS_FOLDER = 'analysis'
MODEL_FOLDER = 'models'
PREVIEW_FOLDER = 'previews'
LOGS_FOLDER = 'logs'

for folder in [UPLOAD_FOLDER, MODEL_FOLDER, PREVIEW_FOLDER, LOGS_FOLDER, ANAL]:
    os.makedirs(folder, exist_ok=True)

ALLOWED_EXTENSIONS = {'csv', 'txt', 'pdf', 'docx', 'png', 'jpg', 'jpeg'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MODEL_FOLDER'] = MODEL_FOLDER
app.config['PREVIEW_FOLDER'] = PREVIEW_FOLDER
app.config['ANALYSIS_FOLDER'] = ANALYSIS_FOLDER

# Load NLP models
nlp = spacy.load("en_core_web_sm")
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

# --- Utility Functions ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def parse_file(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'csv':
        return pd.read_csv(filepath)
    elif ext == 'txt':
        with open(filepath, 'r') as f:
            lines = f.readlines()
        return pd.DataFrame([line.strip().split(',') for line in lines[1:]], columns=lines[0].strip().split(','))
    elif ext == 'docx':
        raise NotImplementedError("Structured DOCX parsing not implemented.")
    elif ext == 'pdf':
        raise NotImplementedError("Structured PDF parsing not implemented.")
    else:
        raise ValueError("Unsupported file type")

def extract_sections_by_style(doc):
    sections = {}
    current_heading = None
    current_content = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_heading:
                sections[current_heading] = ' '.join(current_content).strip()
            current_heading = para.text.strip()
            current_content = []
        elif current_heading:
            current_content.append(para.text.strip())

    if current_heading and current_content:
        sections[current_heading] = ' '.join(current_content).strip()

    return sections
def extract_text_docx(filepath):
    doc = docx.Document(filepath)
    content = []
    for para in doc.paragraphs:
        style = para.style.name
        if para.text.strip():
            content.append((style, para.text.strip()))
    return content
def extract_text_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        if not text.strip():
            # OCR fallback
            images = convert_from_path(filepath)
            for image in images:
                text += pytesseract.image_to_string(image)
        return [("BodyText", text.strip())]
    except Exception as e:
        return [("Error", str(e))]

def extract_text_image(file_path):
    image = Image.open(file_path)
    return pytesseract.image_to_string(image)


def summarize_text(text):
    doc = nlp(text)
    sentences = list(doc.sents)
    summary = ' '.join([str(s) for s in sentences[:3]])
    return summary

def extract_keywords(text):
    doc = nlp(text)
    return list(set(chunk.text.lower() for chunk in doc.noun_chunks if len(chunk.text) > 3))

def extract_named_entities(text):
    doc = nlp(text)
    return [{'text': ent.text, 'label': ent.label_} for ent in doc.ents]

def generate_pdf_report(file_id, extracted, summary=None, keywords=None, entities=None, full_text=''):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Document Summary Report", ln=True, align='C')
    pdf.ln(10)

    if summary:
        pdf.set_font("Arial", size=10, style='B')
        pdf.cell(0, 10, "Summary:", ln=True)
        pdf.set_font("Arial", size=10)
        pdf.multi_cell(0, 10, summary)
        pdf.ln(5)

    if keywords:
        pdf.set_font("Arial", size=10, style='B')
        pdf.cell(0, 10, "Keywords:", ln=True)
        pdf.set_font("Arial", size=10)
        pdf.multi_cell(0, 10, ', '.join(keywords))
        pdf.ln(5)

    if entities:
        pdf.set_font("Arial", size=10, style='B')
        pdf.cell(0, 10, "Named Entities:", ln=True)
        pdf.set_font("Arial", size=10)
        for ent in entities:
            pdf.cell(0, 10, f"{ent['text']} ({ent['label']})", ln=True)
        pdf.ln(5)

    pdf.set_font("Arial", size=10, style='B')
    pdf.cell(0, 10, "Extracted Content:", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 10, full_text)


    for style, content in extracted:
        lines = content.split('\n')
        for line in lines:
            pdf.multi_cell(0, 10, txt=line.strip())

    pdf_path = os.path.join(app.config['ANALYSIS_FOLDER'], f"{file_id}_summary.pdf")
    pdf.output(pdf_path)
    return pdf_path

def generate_pdf_summary(data, output_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Project Analysis Summary", ln=True, align="C")
    pdf.ln(10)

    for heading, content in data.items():
        pdf.set_font("Arial", "B", 14)
        pdf.multi_cell(0, 10, heading)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 10, content)
        pdf.ln(5)

    pdf.output(output_path)

def save_training_log(file_id, original_filename, model_type, accuracy, dataset_shape):
    log_data = {
        "file_id": file_id,
        "original_filename": original_filename,
        "model_type": model_type,
        "accuracy": accuracy,
        "num_rows": dataset_shape[0],
        "num_columns": dataset_shape[1],
        "timestamp": datetime.utcnow().isoformat() + 'Z'
    }

    log_path = os.path.join(LOGS_FOLDER, f"{file_id}_log.json")
    with open(log_path, 'w') as f:
        json.dump(log_data, f, indent=4)

# --- Routes ---
@app.route('/api/train', methods=['POST'])
def train_model():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    model_type = request.form.get('model_type', 'random_forest')

    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_id = str(uuid.uuid4())
        saved_filename = f"{file_id}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)
        file.save(filepath)

        try:
            df = parse_file(filepath)
            if df.shape[0] < 5 or df.shape[1] < 2:
                return jsonify({"error": "Insufficient data for training"}), 400

            preview_path = os.path.join(PREVIEW_FOLDER, f"{file_id}_preview.csv")
            df.head(10).to_csv(preview_path, index=False)

            df = df.apply(pd.to_numeric, errors='ignore')
            X = df.iloc[:, :-1]
            y = df.iloc[:, -1]

            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2)

            if model_type == "random_forest":
                model = RandomForestClassifier()
            elif model_type == "logistic_regression":
                model = LogisticRegression(max_iter=1000)
            else:
                return jsonify({"error": "Unsupported model type"}), 400

            model.fit(X_train, y_train)
            y_pred = model.predict(X_test)
            accuracy = accuracy_score(y_test, y_pred)

            model_path = os.path.join(MODEL_FOLDER, f"{file_id}_{model_type}.pkl")
            joblib.dump(model, model_path)

            save_training_log(file_id, filename, model_type, accuracy, df.shape)

            return jsonify({
                "message": "Model trained successfully",
                "accuracy": accuracy,
                "model_download_url": f"/api/download/model/{file_id}_{model_type}.pkl",
                "preview_download_url": f"/api/download/preview/{file_id}_preview.csv"
            })

        except NotImplementedError as e:
            return jsonify({"error": str(e)}), 501
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    return jsonify({"error": "Unsupported file type"}), 400

@app.route('/api/analyze-doc', methods=['POST'])
def analyze_doc():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_id = str(uuid.uuid4())
        saved_filename = f"{file_id}_{filename}"
        filepath = os.path.join(UPLOAD_FOLDER, saved_filename)
        file.save(filepath)

        try:
            ext = filename.rsplit('.', 1)[1].lower()

            if ext != 'docx':
                return jsonify({"error": "Only .docx supported for heading-style NLP analysis"}), 400

            doc = docx.Document(filepath)
            sections = extract_sections_by_style(doc)

            analysis = {}
            for heading, content in sections.items():
                if len(content) > 30:
                    summary = summarizer(content, max_length=100, min_length=30, do_sample=False)[0]['summary_text']
                else:
                    summary = content
                analysis[heading] = summary

            json_path = os.path.join(LOGS_FOLDER, f"{file_id}_analysis.json")
            with open(json_path, 'w') as f:
                json.dump(analysis, f, indent=4)

            pdf_path = os.path.join(LOGS_FOLDER, f"{file_id}_summary.pdf")
            generate_pdf_summary(analysis, pdf_path)

            return jsonify({
                "message": "Analysis completed",
                "sections": list(analysis.keys()),
                "json_log_url": f"/api/download/log/{file_id}_analysis.json",
                "pdf_summary_url": f"/api/download/log/{file_id}_summary.pdf"
            })

        except Exception as e:
            return jsonify({"error": str(e)}), 500

    return jsonify({"error": "Unsupported file type"}), 400
@app.route('/api/analyze', methods=['POST'])
def analyze_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file type"}), 400

    filename = secure_filename(file.filename)
    file_id = str(uuid.uuid4())
    saved_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{filename}")
    file.save(saved_path)

    ext = filename.rsplit('.', 1)[1].lower()
    if ext == 'docx':
        extracted = extract_text_docx(saved_path)
    elif ext == 'pdf':
        extracted = extract_text_pdf(saved_path)
    elif ext in ['png', 'jpg', 'jpeg']:
        extracted = extract_text_image(saved_path)
    else:
        return jsonify({"error": "Unsupported format"}), 400

    full_text = '\n'.join([text for _, text in extracted])
    summary = summarize_text(full_text)
    keywords = extract_keywords(full_text)
    entities = extract_named_entities(full_text)
    pdf_path = generate_pdf_report(file_id, extracted, summary, keywords, entities)

    analysis_log = {
        "file_id": file_id,
        "filename": filename,
        "summary": summary,
        "keywords": keywords,
        "named_entities": entities,
        "extracted_sections": extracted,
        "report_pdf": f"/api/download/report/{file_id}_summary.pdf",
        "timestamp": datetime.utcnow().isoformat()
    }

    log_path = os.path.join(app.config['ANALYSIS_FOLDER'], f"{file_id}_log.json")
    with open(log_path, 'w') as f:
        json.dump(analysis_log, f, indent=4)

    return jsonify(analysis_log)

@app.route('/api/download/report/<filename>', methods=['GET'])
def download_report(filename):
    return send_from_directory(app.config['ANALYSIS_FOLDER'], filename, as_attachment=True)

@app.route('/api/download/model/<filename>', methods=['GET'])
def download_model(filename):
    return send_from_directory(MODEL_FOLDER, filename, as_attachment=True)

@app.route('/api/download/preview/<filename>', methods=['GET'])
def download_preview(filename):
    return send_from_directory(PREVIEW_FOLDER, filename, as_attachment=True)

@app.route('/api/download/log/<filename>', methods=['GET'])
def download_log(filename):
    return send_from_directory(LOGS_FOLDER, filename, as_attachment=True)

# --- Main ---
if __name__ == '__main__':
    app.run(debug=True)
