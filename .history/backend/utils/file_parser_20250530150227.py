import pandas as pd
import docx
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

def allowed_file(filename, allowed_extensions, ALLOWED EXTENSIONS=None):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def parse_file(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'csv':
        return pd.read_csv(filepath)
    elif ext == 'txt':
        with open(filepath, 'r') as f:
            lines = f.readlines()
        return pd.DataFrame([line.strip().split(',') for line in lines[1:]], columns=lines[0].strip().split(','))
    elif ext == 'docx':
        raise NotImplementedError("Structured DOCX parsing not implemented.")
    elif ext == 'pdf':
        raise NotImplementedError("Structured PDF parsing not implemented.")
    else:
        raise ValueError("Unsupported file type")

def extract_sections_by_style(doc):
    sections = {}
    current_heading = None
    current_content = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_heading:
                sections[current_heading] = ' '.join(current_content).strip()
            current_heading = para.text.strip()
            current_content = []
        elif current_heading:
            current_content.append(para.text.strip())

    if current_heading and current_content:
        sections[current_heading] = ' '.join(current_content).strip()

    return sections

def extract_text_docx(filepath):
    doc = docx.Document(filepath)
    content = []
    for para in doc.paragraphs:
        style = para.style.name
        if para.text.strip():
            content.append((style, para.text.strip()))
    return content

def extract_text_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        if not text.strip():
            # OCR fallback
            images = convert_from_path(filepath)
            for image in images:
                text += pytesseract.image_to_string(image)
        return [("BodyText", text.strip())]
    except Exception as e:
        return [("Error", str(e))]

def extract_text_image(file_path):
    image = Image.open(file_path)
    return pytesseract.image_to_string(image)
