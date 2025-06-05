import pandas as pd
import docx
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pdfplumber
import pytesseract
from PIL import Image


def allowed_file(filename, allowed_extensions):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def parse_file(filepath):
    """Parse structured file types: CSV, TXT, DOCX, PDF."""
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'csv':
        return pd.read_csv(filepath)
    elif ext == 'txt':
        with open(filepath, 'r') as f:
            lines = f.readlines()
        if not lines:
            raise ValueError("TXT file is empty")
        return pd.DataFrame(
            [line.strip().split(',') for line in lines[1:] if line.strip()],
            columns=lines[0].strip().split(',')
        )
    elif ext == 'docx':
        return extract_tables_docx(filepath)
    elif ext == 'pdf':
        return extract_tables_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

def extract_sections_by_style(doc):
    """Extract DOCX text sections grouped by headings."""
    sections = {}
    current_heading = None
    current_content = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_heading:
                sections[current_heading] = ' '.join(current_content).strip()
            current_heading = para.text.strip()
            current_content = []
        elif current_heading:
            current_content.append(para.text.strip())

    if current_heading and current_content:
        sections[current_heading] = ' '.join(current_content).strip()

    return sections

def extract_text_docx(filepath):
    """Extract all text from a DOCX file with style labels."""
    try:
        doc = docx.Document(filepath)
        return [
            (para.style.name, para.text.strip())
            for para in doc.paragraphs if para.text.strip()
        ]
    except Exception as e:
        return [("Error", f"Failed to read DOCX: {e}")]

def extract_text_pdf(filepath):
    """Extract text from a PDF using PyPDF2 or OCR fallback."""
    try:
        reader = PdfReader(filepath)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        if not text.strip():
            # OCR fallback
            images = convert_from_path(filepath)
            for image in images:
                text += pytesseract.image_to_string(image)
        return [("BodyText", text.strip())]
    except Exception as e:
        return [("Error", f"Failed to extract PDF text: {e}")]

def extract_text_image(filepath):
    """Extract text from an image using OCR."""
    try:
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image)
        return [("BodyText", text.strip())]
    except Exception as e:
        return [("Error", f"Failed to extract image text: {e}")]

def extract_tables_docx(filepath):
    """Extract tables from a DOCX file and return as list of DataFrames."""
    try:
        doc = docx.Document(filepath)
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                tables.append(df)
        return tables if tables else [{"message": "No tables found in DOCX."}]
    except Exception as e:
        return [{"error": f"Failed to extract tables from DOCX: {e}"}]

def extract_tables_pdf(filepath):
    """Extract tables from a PDF using pdfplumber."""
    try:
        tables = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                extracted = page.extract_tables()
                for table in extracted:
                    if table:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        tables.append(df)
        return tables if tables else [{"message": "No tables found in PDF."}]
    except Exception as e:
        return [{"error": f"Failed to extract tables from PDF: {e}"}]
